#!/usr/bin/env python3
"""Build a disposable 60k+ character DOCX for Vera Word Host E2E testing.

The fixture deliberately contains only invented contract text.  Its paragraph
boundaries exercise the add-in's long-document segmentation and its two unique
clauses make it possible to prove that late-document suggestions still point to
the right source.  It is a test input, not a user-facing legal work product.
"""

from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_OUTPUT = Path("/tmp/vera-word-60k-host-e2e.docx")
ENGLISH_TARGET = (
    "The Supplier may change the Service Fees at any time without prior notice."
)
CHINESE_TARGET = "供应商可在未经客户书面同意的情况下随时变更处理目的。"
MAX_PARAGRAPH_CHARS = 16_000


def set_run_font(run, western: str = "Calibri", east_asia: str = "Noto Sans CJK SC"):
    run.font.name = western
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "en-US")
    language.set(qn("w:eastAsia"), "zh-CN")
    run._element.rPr.append(language)


def configure_document(doc: Document):
    """Apply the standard_business_brief geometry and type tokens."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_body(doc: Document, text: str):
    if len(text) >= MAX_PARAGRAPH_CHARS:
        raise ValueError(f"Paragraph exceeds the {MAX_PARAGRAPH_CHARS:,}-character guardrail")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    if any(ord(character) > 127 for character in text):
        # LibreOffice's renderer only reliably honors this fixture's CJK font
        # when it is also the western/HAnsi font on a mixed-script run.
        set_run_font(run, western="Noto Sans CJK SC", east_asia="Noto Sans CJK SC")
    else:
        set_run_font(run)
    run.font.size = Pt(11)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_heading(text, level=level)
    if any(ord(character) > 127 for character in text):
        for run in paragraph.runs:
            set_run_font(run, western="Noto Sans CJK SC", east_asia="Noto Sans CJK SC")
    return paragraph


def filler_clause(index: int) -> str:
    """Return a numbered, non-repeating invented clause of roughly 1,200 characters."""
    clause_id = f"SYN-{index:03d}"
    sentence = (
        f"Section {clause_id} concerns an entirely fictional service arrangement for host-test "
        "purposes. The parties shall record operational requests in a shared test ledger, use "
        "reasonable efforts to maintain continuity, and give written notice of a material issue "
        "within ten business days. Neither party may rely on this synthetic language outside the "
        "fixture. The Supplier will maintain a documented release process, while the Customer will "
        "designate an authorized test contact for non-production questions. Any disagreement will "
        "be discussed by the nominated contacts before either side starts a formal escalation. "
        "All timelines in this paragraph are illustrative only and do not describe a real product, "
        "person, account, client, or transaction. "
    )
    return (sentence * 2) + f"End of synthetic clause {clause_id}."


def add_target_paragraphs(doc: Document):
    add_heading(doc, "Fee adjustment review target", level=2)
    add_body(
        doc,
        f"Special review target. {ENGLISH_TARGET} "
        "This sentence appears once in the fixture and is intentionally positioned after the "
        "first twenty thousand characters so a host run can verify late-document source matching.",
    )


def add_chinese_target_paragraph(doc: Document):
    add_heading(doc, "中文处理目的审阅目标", level=2)
    add_body(
        doc,
        f"特别审阅目标：{CHINESE_TARGET}"
        "该句在夹具中仅出现一次，并被放置在四万五千字符之后，以验证中文来源锚点不会被前文截断。",
    )


def text_metrics(doc: Document) -> tuple[str, int, int, int, int]:
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    english_offset = text.index(ENGLISH_TARGET)
    chinese_offset = text.index(CHINESE_TARGET)
    return text, len(doc.paragraphs), english_offset, chinese_offset, max(
        len(paragraph.text) for paragraph in doc.paragraphs
    )


def validate_output(output: Path):
    """Check the package and the exact text metrics a Host run relies on."""
    if not zipfile.is_zipfile(output):
        raise ValueError("Generated output is not a valid DOCX ZIP package")
    with zipfile.ZipFile(output) as package:
        names = set(package.namelist())
        required = {"[Content_Types].xml", "word/document.xml"}
        missing = required - names
        if missing:
            raise ValueError(f"DOCX package is missing required parts: {sorted(missing)}")
        if package.testzip() is not None:
            raise ValueError("DOCX ZIP integrity check failed")

    reopened = Document(output)
    text, paragraph_count, english_offset, chinese_offset, longest = text_metrics(reopened)
    errors = []
    if len(text) <= 60_000:
        errors.append("document text must exceed 60,000 characters")
    if english_offset <= 20_000:
        errors.append("English target must appear after 20,000 characters")
    if chinese_offset <= 45_000:
        errors.append("Chinese target must appear after 45,000 characters")
    if text.count(ENGLISH_TARGET) != 1 or text.count(CHINESE_TARGET) != 1:
        errors.append("each target must appear exactly once")
    if longest >= MAX_PARAGRAPH_CHARS:
        errors.append("each paragraph must be shorter than 16,000 characters")
    if (
        len(reopened.paragraphs) < 3
        or reopened.paragraphs[0].text.strip()
        or reopened.paragraphs[1].text.strip()
        or not reopened.paragraphs[2].text.strip()
    ):
        errors.append("fixture must preserve exactly two leading empty paragraphs")
    if errors:
        raise ValueError("; ".join(errors))

    return {
        "characters": len(text),
        "paragraphs": paragraph_count,
        "english_offset": english_offset,
        "chinese_offset": chinese_offset,
        "longest_paragraph": longest,
    }


def build(output: Path):
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # python-docx starts a new document without body paragraphs. Create the
    # first slot, then reuse ``doc.paragraphs[0]`` and add exactly one more so
    # the first content is paragraph three. The line breaks prevent Word from
    # coalescing either blank slot.
    doc.add_paragraph()
    doc.paragraphs[0].add_run().add_break()
    doc.add_paragraph().add_run().add_break()
    add_heading(doc, "Vera Word 60k Host E2E Fixture", level=1)
    note = add_body(
        doc,
        "Synthetic test content only. This disposable fixture contains no client, personal, or confidential information.",
    )
    note.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_heading(doc, "Synthetic agreement body", level=2)
    for index in range(1, 25):
        add_body(doc, filler_clause(index))

    add_target_paragraphs(doc)

    add_heading(doc, "Late-document continuation", level=2)
    for index in range(25, 33):
        add_body(doc, filler_clause(index))

    add_chinese_target_paragraph(doc)

    add_heading(doc, "Closing synthetic provisions", level=2)
    for index in range(33, 50):
        add_body(doc, filler_clause(index))

    doc.core_properties.title = "Vera Word 60k Host E2E Fixture"
    doc.core_properties.subject = "Synthetic long-document Office Add-in validation"
    doc.core_properties.author = "Vera QA"
    doc.core_properties.keywords = "synthetic, office add-in, word, host, long document"
    doc.save(output)
    return validate_output(output)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    metrics = build(args.output)
    print(f"output={args.output}")
    print(f"characters={metrics['characters']}")
    print(f"paragraphs={metrics['paragraphs']}")
    print(f"english_target_offset={metrics['english_offset']}")
    print(f"chinese_target_offset={metrics['chinese_offset']}")
    print(f"longest_paragraph={metrics['longest_paragraph']}")


if __name__ == "__main__":
    main()
